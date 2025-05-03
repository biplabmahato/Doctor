import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import base64
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Gemini
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Initialize the Gemini Pro model
model = genai.GenerativeModel('gemini-1.5-flash')   

# Document generation functions (same as before)
def generate_docx(result):
    doc = Document()
    doc.add_heading('Comprehensive Healthcare Report', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Full Medical Report</a>'

# Streamlit UI (same as before)
st.set_page_config(layout="wide")
st.title("AI-Powered Medical Diagnosis System ")

# Patient Information Section
st.header("Patient Information")
col1, col2 = st.columns(2)
with col1:
    gender = st.selectbox('Gender', ('Male', 'Female', 'Other'))
    age = st.number_input('Age', min_value=0, max_value=120, value=25)
    height = st.number_input('Height (cm)', min_value=50, max_value=250, value=170)
with col2:
    weight = st.number_input('Weight (kg)', min_value=10, max_value=300, value=70)
    blood_type = st.selectbox('Blood Type', ('A+', 'A-', 'B+', 'B-', 'AB+', 'AB-', 'O+', 'O-', 'Unknown'))
    allergies = st.text_input('Known Allergies', 'None')

# Symptoms and History Section
st.header("Symptoms and Medical History")
symptoms = st.text_area('Current Symptoms', 'e.g., fever for 3 days, sharp pain in lower right abdomen')
duration = st.selectbox('Duration of Symptoms', 
                       ('Less than 24 hours', '1-3 days', '3-7 days', '1-2 weeks', 'More than 2 weeks'))
medical_history = st.text_area('Medical History', 'e.g., diabetes type 2 diagnosed 2018, hypertension')
current_medications = st.text_area('Current Medications', 'e.g., metformin 500mg twice daily, lisinopril 10mg daily')

# Generate Report Button
if st.button("Generate Medical Report"):
    # Construct the prompt
    prompt = f"""
    Act as an expert medical diagnostician. Analyze this patient case:

    Patient Profile:
    - Gender: {gender}
    - Age: {age}
    - Height: {height} cm
    - Weight: {weight} kg
    - Blood Type: {blood_type}
    - Allergies: {allergies}

    Symptoms:
    - {symptoms} (Duration: {duration})

    Medical Background:
    - History: {medical_history}
    - Current Medications: {current_medications}

    Provide a comprehensive report with:
    1. Differential diagnosis (list 3 most likely conditions)
    2. Recommended diagnostic tests
    3. Treatment plan (medications + lifestyle)
    4. Follow-up recommendations
    5. Red flags to watch for

    Format the response in clear markdown with headings.
    """

    with st.spinner('Generating report ...'):
        try:
            # Generate content
            response = model.generate_content(prompt)
            
            # Display and make downloadable
            st.markdown(response.text)
            docx_file = generate_docx(response.text)
            st.markdown(get_download_link(docx_file, "medical_report.docx"), unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"Error generating report: {str(e)}")
